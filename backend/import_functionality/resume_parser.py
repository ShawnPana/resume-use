import json
import os
import sys
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime

import anthropic
from PyPDF2 import PdfReader
from docx import Document
from dotenv import load_dotenv

# Load .env from backend directory
env_path = Path(__file__).parent.parent / '.env'
load_dotenv(env_path)


class ResumeParser:
    def __init__(self, api_key: Optional[str] = None):
        """
        Initialize the ResumeParser with Anthropic API key.

        Args:
            api_key: Anthropic API key. If not provided, will look for ANTHROPIC_API_KEY env variable.
        """
        self.api_key = api_key or os.getenv("ANTHROPIC_API_KEY")
        if not self.api_key:
            raise ValueError("Anthropic API key is required. Set ANTHROPIC_API_KEY environment variable or pass it as parameter.")

        self.client = anthropic.Anthropic(api_key=self.api_key)

        # Load the template structure
        template_path = Path(__file__).parent.parent / "data" / "import_template.json"
        with open(template_path, 'r') as f:
            self.template = json.load(f)

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from a PDF file."""
        try:
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"Error reading PDF file: {str(e)}")

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file."""
        try:
            doc = Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)

            return "\n".join(text)
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {str(e)}")

    def extract_text(self, file_path: str) -> str:
        """Extract text from either PDF or DOCX file."""
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = file_path.suffix.lower()

        if extension == '.pdf':
            return self.extract_text_from_pdf(str(file_path))
        elif extension in ['.docx', '.doc']:
            return self.extract_text_from_docx(str(file_path))
        else:
            raise ValueError(f"Unsupported file type: {extension}. Only PDF and DOCX files are supported.")

    def parse_with_anthropic(self, resume_text: str) -> Dict[str, Any]:
        """
        Use Anthropic's Claude API to parse resume text into the structured format.
        """
        system_prompt = """You are a resume parser that extracts information from resumes and formats it into a specific JSON structure.

IMPORTANT: Return ONLY valid JSON without any markdown code blocks or additional text.

Extract the following information from the resume and format it according to this exact structure:

{
  "header": {
    "email": "email address",
    "github": "github profile URL (full URL)",
    "lastUpdated": "",  // Leave empty
    "linkedin": "linkedin profile URL (full URL)",
    "location": "city, state or city, country",
    "name": "full name",
    "phone": "phone number",
    "skills": {
      "languages": ["programming languages"],
      "webDevelopment": ["web frameworks and tools"],
      "aiML": ["AI/ML tools and frameworks"],
      "cloudData": ["cloud and database technologies"],
      "tools": ["other development tools"]
    },
    "tagline": "professional tagline or objective",
    "website": "personal website URL"
  },
  "education": {
    "coursework": ["list of relevant courses"],
    "degree": "degree type (e.g., Bachelor of Science (B.S.))",
    "endDate": "MM/YYYY or 'Present'",
    "gpa": "GPA value",
    "major": "field of study",
    "startDate": "MM/YYYY",
    "university": "university name"
  },
  "experience": [
    {
      "description": "brief description of role and achievements",
      "endDate": "MM/YYYY or 'Present'",
      "position": "job title",
      "startDate": "MM/YYYY",
      "title": "company name",
      "url": "company website if mentioned"
    }
  ],
  "projects": [
    {
      "award": "any awards won",
      "date": "MM/YYYY",
      "description": "project description",
      "endDate": "MM/YYYY or 'Present' if ongoing",
      "event": "hackathon or event name if applicable",
      "organization": "organization if applicable",
      "title": "project name",
      "url": "project URL if available"
    }
  ]
}

Guidelines:
- If information is not available, use empty string "" for strings, empty array [] for arrays
- For skills, intelligently categorize them into the appropriate subcategories
- Format all dates as MM/YYYY (e.g., "09/2024")
- If only year is provided, use "01/YYYY" for start dates and "12/YYYY" for end dates
- For current positions/projects, use "Present" as endDate
- Extract as much relevant information as possible from the resume
- Combine multiple bullet points into a single description field with clear, concise text
- Do not invent information that is not in the resume"""

        user_prompt = f"Parse the following resume and extract information into the JSON format:\n\n{resume_text}"

        try:
            response = self.client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=4000,
                temperature=0,
                system=system_prompt,
                messages=[
                    {"role": "user", "content": user_prompt}
                ]
            )

            # Extract the JSON from the response
            response_text = response.content[0].text.strip()

            # Try to parse the JSON
            try:
                parsed_data = json.loads(response_text)
                return self.validate_and_clean_data(parsed_data)
            except json.JSONDecodeError:
                # If direct parsing fails, try to extract JSON from the response
                import re
                json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
                if json_match:
                    parsed_data = json.loads(json_match.group())
                    return self.validate_and_clean_data(parsed_data)
                else:
                    raise ValueError("Could not extract valid JSON from API response")

        except Exception as e:
            raise Exception(f"Error calling Anthropic API: {str(e)}")

    def validate_and_clean_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Validate and clean the parsed data to match the expected format.
        """
        # Ensure all required top-level keys exist
        if 'header' not in data:
            data['header'] = {}
        if 'education' not in data:
            data['education'] = {}
        if 'experience' not in data:
            data['experience'] = []
        if 'projects' not in data:
            data['projects'] = []

        # Clean header
        header_fields = ['email', 'github', 'lastUpdated', 'linkedin', 'location',
                        'name', 'phone', 'skills', 'tagline', 'website']
        for field in header_fields:
            if field not in data['header']:
                if field == 'skills':
                    data['header']['skills'] = {}
                else:
                    data['header'][field] = ""

        # Ensure skills is properly formatted as a dict
        if not isinstance(data['header']['skills'], dict):
            data['header']['skills'] = {}

        # Clean education
        edu_fields = ['coursework', 'degree', 'endDate', 'gpa', 'major', 'startDate', 'university']
        for field in edu_fields:
            if field not in data['education']:
                if field == 'coursework':
                    data['education'][field] = []
                else:
                    data['education'][field] = ""

        # Ensure coursework is a list
        if not isinstance(data['education'].get('coursework'), list):
            data['education']['coursework'] = []

        # Clean experience entries
        exp_fields = ['description', 'endDate', 'position', 'startDate', 'title', 'url']
        for i, exp in enumerate(data['experience']):
            for field in exp_fields:
                if field not in exp:
                    data['experience'][i][field] = ""

        # Clean project entries
        proj_fields = ['award', 'date', 'description', 'endDate', 'event', 'organization', 'title', 'url']
        for i, proj in enumerate(data['projects']):
            for field in proj_fields:
                if field not in proj:
                    data['projects'][i][field] = ""

        # Set lastUpdated to current date
        data['header']['lastUpdated'] = datetime.now().strftime("%m/%Y")

        return data

    def parse_resume(self, file_path: str) -> Dict[str, Any]:
        """
        Main method to parse a resume file and return structured data.

        Args:
            file_path: Path to the resume file (PDF or DOCX)

        Returns:
            Dictionary containing parsed resume data in the template format
        """
        # Extract text from file
        resume_text = self.extract_text(file_path)

        if not resume_text or len(resume_text.strip()) < 50:
            raise ValueError("Could not extract sufficient text from the resume file")

        # Parse with Anthropic API
        parsed_data = self.parse_with_anthropic(resume_text)

        return parsed_data

    def save_to_file(self, data: Dict[str, Any], output_path: Optional[str] = None) -> str:
        """
        Save parsed data to a JSON file.

        Args:
            data: Parsed resume data
            output_path: Optional path for output file. If not provided, will save to data directory.

        Returns:
            Path to the saved file
        """
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = Path(__file__).parent.parent / "data" / f"parsed_resume_{timestamp}.json"

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        with open(output_path, 'w') as f:
            json.dump(data, f, indent=2)

        return str(output_path)


def main():
    """Command-line interface for the resume parser."""
    if len(sys.argv) < 2:
        print("Usage: python resume_parser.py <resume_file_path> [output_json_path]")
        print("\nExample:")
        print("  python resume_parser.py resume.pdf")
        print("  python resume_parser.py resume.docx output.json")
        sys.exit(1)

    resume_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None

    try:
        # Initialize parser
        parser = ResumeParser()

        print(f"Parsing resume: {resume_file}")

        # Parse the resume
        parsed_data = parser.parse_resume(resume_file)

        # Save to file
        saved_path = parser.save_to_file(parsed_data, output_file)

        print(f"Successfully parsed resume!")
        print(f"Output saved to: {saved_path}")

        # Print summary
        if parsed_data.get('header', {}).get('name'):
            print(f"\nExtracted information for: {parsed_data['header']['name']}")
        if parsed_data.get('experience'):
            print(f"Found {len(parsed_data['experience'])} experience entries")
        if parsed_data.get('projects'):
            print(f"Found {len(parsed_data['projects'])} project entries")

    except Exception as e:
        print(f"Error: {str(e)}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()